"""
File processing utilities for extracting text content from various file types.
This content will be used for AI context in chat conversations.
"""

from pathlib import Path
from typing import Optional, Tuple

# Text extraction libraries
try:
    import PyPDF2
    HAS_PDF = True
except ImportError:
    HAS_PDF = False

try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    import pandas as pd
    HAS_PANDAS = True
except ImportError:
    HAS_PANDAS = False


async def extract_text_from_file(file_path: str, file_type: str) -> Tuple[bool, Optional[str], Optional[str]]:
    """
    Extract text content from a file.
    
    Args:
        file_path: Path to the file
        file_type: File extension (without dot)
        
    Returns:
        Tuple of (success, extracted_text, error_message)
    """
    try:
        file_path_obj = Path(file_path)
        
        if not file_path_obj.exists():
            return False, None, "File not found"
        
        if file_type.lower() in ['txt', 'md', 'py', 'js', 'json', 'yaml', 'yml', 'xml', 'html', 'css']:
            # Plain text files
            return await _extract_text_file(file_path_obj)
        
        elif file_type.lower() == 'pdf':
            return await _extract_pdf_text(file_path_obj)
        
        elif file_type.lower() in ['docx', 'doc']:
            return await _extract_docx_text(file_path_obj)
        
        elif file_type.lower() in ['xlsx', 'xls', 'csv']:
            return await _extract_spreadsheet_text(file_path_obj, file_type)
        
        else:
            return False, None, f"Unsupported file type: {file_type}"
            
    except Exception as e:
        return False, None, f"Error processing file: {str(e)}"


async def _extract_text_file(file_path: Path) -> Tuple[bool, Optional[str], Optional[str]]:
    """Extract text from plain text files."""
    try:
        # Try different encodings
        encodings = ['utf-8', 'utf-8-sig', 'latin-1', 'cp1252']
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    content = f.read()
                    return True, content, None
            except UnicodeDecodeError:
                continue
        
        return False, None, "Could not decode file with any common encoding"
        
    except Exception as e:
        return False, None, f"Error reading text file: {str(e)}"


async def _extract_pdf_text(file_path: Path) -> Tuple[bool, Optional[str], Optional[str]]:
    """Extract text from PDF files."""
    if not HAS_PDF:
        return False, None, "PyPDF2 not installed. Cannot process PDF files."
    
    try:
        text_content = ""
        
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                text_content += page.extract_text() + "\n"
        
        if text_content.strip():
            return True, text_content, None
        else:
            return False, None, "No text content found in PDF"
            
    except Exception as e:
        return False, None, f"Error reading PDF: {str(e)}"


async def _extract_docx_text(file_path: Path) -> Tuple[bool, Optional[str], Optional[str]]:
    """Extract text from DOCX files."""
    if not HAS_DOCX:
        return False, None, "python-docx not installed. Cannot process DOCX files."
    
    try:
        doc = Document(file_path)
        text_content = ""
        
        for paragraph in doc.paragraphs:
            text_content += paragraph.text + "\n"
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text_content += cell.text + "\t"
                text_content += "\n"
        
        if text_content.strip():
            return True, text_content, None
        else:
            return False, None, "No text content found in DOCX"
            
    except Exception as e:
        return False, None, f"Error reading DOCX: {str(e)}"


async def _extract_spreadsheet_text(file_path: Path, file_type: str) -> Tuple[bool, Optional[str], Optional[str]]:
    """Extract text from spreadsheet files."""
    if not HAS_PANDAS:
        return False, None, "pandas not installed. Cannot process spreadsheet files."
    
    try:
        if file_type.lower() == 'csv':
            df = pd.read_csv(file_path)
        else:
            df = pd.read_excel(file_path)
        
        # Convert dataframe to text representation
        text_content = f"Spreadsheet content ({file_type.upper()}):\n\n"
        text_content += df.to_string(index=False)
        
        return True, text_content, None
        
    except Exception as e:
        return False, None, f"Error reading spreadsheet: {str(e)}"


def get_processing_requirements() -> dict:
    """Get information about which file types can be processed."""
    return {
        "text_files": ["txt", "md", "py", "js", "json", "yaml", "yml", "xml", "html", "css"],
        "pdf_files": ["pdf"] if HAS_PDF else [],
        "word_files": ["docx", "doc"] if HAS_DOCX else [],
        "spreadsheet_files": ["xlsx", "xls", "csv"] if HAS_PANDAS else [],
        "missing_dependencies": {
            "PyPDF2": not HAS_PDF,
            "python-docx": not HAS_DOCX, 
            "pandas": not HAS_PANDAS
        }
    }
